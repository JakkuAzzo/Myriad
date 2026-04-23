from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


def insert_paragraph_after(document: Document, paragraph: Paragraph, text: str = '', style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    if text:
      created.add_run(text)
    if style_name and style_name in document.styles:
      created.style = document.styles[style_name]
    return created


def find_anchor_paragraph(document: Document, keywords: Iterable[str]) -> Paragraph:
    lowered = [k.lower() for k in keywords]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip().lower()
        if not text:
            continue
        if any(keyword in text for keyword in lowered):
            return paragraph
    raise RuntimeError(f'Could not find anchor paragraph for keywords: {keywords}')


def read_text(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(path)
    return path.read_text(encoding='utf-8')


def extract_h2_section(markdown: str, heading: str) -> str:
    lines = markdown.splitlines()
    start = None
    for i, line in enumerate(lines):
        if line.strip() == f'## {heading}':
            start = i + 1
            break

    if start is None:
        raise RuntimeError(f'Section not found: {heading}')

    end = len(lines)
    for i in range(start, len(lines)):
        if lines[i].startswith('## '):
            end = i
            break

    section = '\n'.join(lines[start:end]).strip('\n')
    return section


def section_already_present(document: Document, unique_probe_text: str) -> bool:
    probe = unique_probe_text.strip().lower()
    if not probe:
        return False
    for paragraph in document.paragraphs:
        if paragraph.text.strip().lower() == probe:
            return True
    return False


def inject_markdown_block(document: Document, anchor: Paragraph, markdown_block: str) -> Paragraph:
    current = anchor
    in_code_block = False

    for raw_line in markdown_block.splitlines():
        line = raw_line.rstrip('\n')
        stripped = line.strip()

        if stripped == '```':
            in_code_block = not in_code_block
            continue

        if stripped == '```mermaid':
            in_code_block = True
            continue

        if in_code_block:
            current = insert_paragraph_after(document, current, line)
            continue

        if not stripped:
            current = insert_paragraph_after(document, current, '')
            continue

        if stripped.startswith('### '):
            current = insert_paragraph_after(document, current, stripped[4:].strip(), 'Heading 3')
            continue

        if stripped.startswith('#### '):
            style = 'Heading 4' if 'Heading 4' in document.styles else None
            current = insert_paragraph_after(document, current, stripped[5:].strip(), style)
            continue

        if stripped.startswith('- '):
            style = 'List Bullet' if 'List Bullet' in document.styles else None
            current = insert_paragraph_after(document, current, stripped[2:].strip(), style)
            continue

        if _is_numbered_list_line(stripped):
            dot_idx = stripped.find('.')
            text = stripped[dot_idx + 1 :].strip()
            style = 'List Number' if 'List Number' in document.styles else None
            current = insert_paragraph_after(document, current, text, style)
            continue

        current = insert_paragraph_after(document, current, stripped)

    return current


def _is_numbered_list_line(text: str) -> bool:
    dot_idx = text.find('.')
    if dot_idx <= 0:
        return False
    prefix = text[:dot_idx]
    return prefix.isdigit() and len(text) > dot_idx + 1 and text[dot_idx + 1] == ' '


def inject_figures_from_manifest(
    document: Document,
    anchor: Paragraph,
    manifest_path: Path,
    image_width_in: float,
) -> Paragraph:
    manifest = json.loads(read_text(manifest_path))
    current = anchor

    if not isinstance(manifest, list):
        raise RuntimeError(f'Invalid manifest format: {manifest_path}')

    current = insert_paragraph_after(document, current, 'Appendix E Figure Plates', 'Heading 3')
    current = insert_paragraph_after(
        document,
        current,
        'Rendered diagrams generated from Mermaid artefacts for direct dissertation insertion.',
    )

    for item in manifest:
        image_path = Path(item['imageFile'])
        caption = item.get('caption', image_path.stem)

        if not image_path.exists():
            raise FileNotFoundError(f'Missing rendered image: {image_path}')

        pic_paragraph = insert_paragraph_after(document, current, '')
        run = pic_paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(image_width_in))

        current = insert_paragraph_after(document, pic_paragraph, caption)

    return current


def main() -> None:
    parser = argparse.ArgumentParser(description='Inject dissertation sections into a .docx by heading anchors.')
    parser.add_argument('--docx', required=True, help='Absolute path to dissertation .docx file')
    parser.add_argument(
        '--sections-markdown',
        default=None,
        help='Markdown file containing Chapter 2/3/4 and Appendices sections. Defaults to artifacts/dissertation/depth-evidence-pack.md',
    )
    parser.add_argument(
        '--manifest',
        default=None,
        help='Optional path to rendered Mermaid figures manifest.json',
    )
    parser.add_argument('--output', default=None, help='Optional output path. Defaults to overwrite input docx.')
    parser.add_argument('--image-width', type=float, default=5.8, help='Image width in inches for inserted figures')
    args = parser.parse_args()

    docx_path = Path(args.docx).expanduser().resolve()
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    default_sections = Path(__file__).resolve().parent.parent / 'artifacts' / 'dissertation' / 'depth-evidence-pack.md'
    sections_path = Path(args.sections_markdown).expanduser().resolve() if args.sections_markdown else default_sections

    markdown = read_text(sections_path)
    document = Document(str(docx_path))

    jobs = [
        {
            'name': 'Chapter 2 block',
            'section': 'Chapter 2 Addition: Personas, User Stories, and Use Cases',
            'anchor_keywords': ['chapter 2', 'analysis and design', 'analysis'],
            'probe': '2.x Persona-Driven Requirements Refinement',
        },
        {
            'name': 'Chapter 3 block',
            'section': 'Chapter 3 Addition: Database, API, and Architecture Detail',
            'anchor_keywords': ['chapter 3', 'realisation', 'implementation'],
            'probe': '3.x Data Model and Persistence Design',
        },
        {
            'name': 'Chapter 4 block',
            'section': 'Chapter 4 Addition: Full Usability and System Test Evidence',
            'anchor_keywords': ['chapter 4', 'validation', 'testing'],
            'probe': '4.x Usability Protocol and Participant Data',
        },
        {
            'name': 'Appendix pointers',
            'section': 'Appendices Additions',
            'anchor_keywords': ['appendices', 'appendix'],
            'probe': 'Appendix A: Interview Prompt Set',
        },
    ]

    for job in jobs:
        if section_already_present(document, job['probe']):
            print(f"SKIP_ALREADY_PRESENT {job['name']}")
            continue

        section_text = extract_h2_section(markdown, job['section'])
        anchor = find_anchor_paragraph(document, job['anchor_keywords'])
        inject_markdown_block(document, anchor, section_text)
        print(f"INJECTED {job['name']}")

    if args.manifest:
        manifest_path = Path(args.manifest).expanduser().resolve()
        if not manifest_path.exists():
            raise FileNotFoundError(manifest_path)

        if not section_already_present(document, 'Appendix E Figure Plates'):
            appendix_anchor = find_anchor_paragraph(document, ['appendix e', 'appendices', 'appendix'])
            inject_figures_from_manifest(document, appendix_anchor, manifest_path, args.image_width)
            print('INJECTED Appendix E figures')
        else:
            print('SKIP_ALREADY_PRESENT Appendix E figures')

    output_path = Path(args.output).expanduser().resolve() if args.output else docx_path
    document.save(str(output_path))
    print(f'SAVED {output_path}')


if __name__ == '__main__':
    main()
