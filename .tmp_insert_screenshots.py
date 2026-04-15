from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches

DOC_PATH = Path('/Users/nathanbrown-bennett/Said Osman/Myriad_Dissertation_Said_Osman_Completed.docx')
IMG_DIR = Path('/Users/nathanbrown-bennett/Said Osman/artifacts/dissertation')

images = [
    (IMG_DIR / 'myriad_home.png', 'Figure 5: Myriad home screen and primary controls.'),
    (IMG_DIR / 'myriad_stats.png', 'Figure 6: Dashboard statistics overview for behavioural analytics output.'),
    (IMG_DIR / 'poc_category_usage.png', 'Figure 7: Category-usage visualisation from local processed events.'),
    (IMG_DIR / 'poc_sentiment_trend.png', 'Figure 8: Sentiment-trend visualisation used for reflective feedback.'),
    (IMG_DIR / 'poc_test_outcomes.png', 'Figure 9: Validation and test outcomes snapshot from the artefact.'),
]

for path, _ in images:
    if not path.exists():
        raise FileNotFoundError(f'Missing image: {path}')

doc = Document(DOC_PATH)


def insert_paragraph_after(paragraph, text='', style_name=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    created.text = text
    if style_name:
        created.style = doc.styles[style_name]
    return created

# Find insertion point under Realisation chapter after 3.3.3
anchor = None
for p in doc.paragraphs:
    if p.text.strip() == '3.3.3 Example Output Structure':
        anchor = p
        break

if anchor is None:
    raise RuntimeError('Could not find insertion anchor heading 3.3.3 Example Output Structure')

cur = anchor
cur = insert_paragraph_after(cur, '3.3.4 Application Screenshots', 'Heading 3')
cur = insert_paragraph_after(cur, 'The following screenshots provide direct visual evidence of the implemented interface and analytics outputs referenced in this chapter.', None)

for image_path, caption in images:
    pic_p = insert_paragraph_after(cur, '', None)
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))
    cap_p = insert_paragraph_after(pic_p, caption, None)
    cur = cap_p

# Update manual Table of Figures section with matching entries
# Find heading then append new lines before next heading
tof_idx = None
paras = list(doc.paragraphs)
for i, p in enumerate(paras):
    if p.text.strip() == 'Table of Figures' and p.style and p.style.name == 'Heading 1':
        tof_idx = i
        break

if tof_idx is not None:
    next_heading_idx = len(paras)
    for j in range(tof_idx + 1, len(paras)):
        pj = paras[j]
        if pj.style and pj.style.name.startswith('Heading'):
            next_heading_idx = j
            break

    existing_texts = {pp.text.strip() for pp in doc.paragraphs}
    append_after = doc.paragraphs[tof_idx]
    for line in [
        'Figure 5: Myriad home screen and primary controls',
        'Figure 6: Dashboard statistics overview for behavioural analytics output',
        'Figure 7: Category-usage visualisation from local processed events',
        'Figure 8: Sentiment-trend visualisation used for reflective feedback',
        'Figure 9: Validation and test outcomes snapshot from the artefact',
    ]:
        if line not in existing_texts:
            append_after = insert_paragraph_after(append_after, line, None)

# Save

doc.save(DOC_PATH)
print(f'Inserted screenshots into: {DOC_PATH}')
