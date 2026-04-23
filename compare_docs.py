import docx
import os

def process_doc(filename):
    if not os.path.exists(filename):
        print(f"File {filename} not found.")
        return
    
    print(f"\n--- Processing: {filename} ---")
    doc = docx.Document(filename)
    
    print("\nHeadings found:")
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            print(f"Index {i}: [{para.style.name}] {para.text[:100]}")
            
    print("\nFirst 40 non-empty paragraphs:")
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"Index {i}: {text[:100]}")
            count += 1
        if count >= 40:
            break

process_doc('Myriad_Dissertation_Said_Osman_Completed.docx')
process_doc('Myriad_Dissertation_Said_Osman_Completed_injected.docx')
