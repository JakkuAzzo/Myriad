from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


def insert_paragraph_after(document: Document, paragraph: Paragraph, text: str = '', style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    created = Paragraph(new_p, paragraph._parent)
    if text:
        created.add_run(text)
    if style_name and style_name in document.styles:
        created.style = document.styles[style_name]
    return created


def find_paragraph_by_exact_text(document: Document, text: str) -> Paragraph:
    needle = text.strip()
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == needle:
            return paragraph
    raise RuntimeError(f'Could not find paragraph: {text}')


def get_paragraph_before(document: Document, heading_text: str) -> Paragraph:
    paragraphs = list(document.paragraphs)
    heading_idx = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == heading_text.strip():
            heading_idx = idx
            break

    if heading_idx is None:
        raise RuntimeError(f'Could not find heading: {heading_text}')
    if heading_idx == 0:
        raise RuntimeError(f'Heading has no previous paragraph: {heading_text}')

    return paragraphs[heading_idx - 1]


def section_present(document: Document, probe: str) -> bool:
    target = probe.strip().lower()
    if not target:
        return False
    return any(p.text.strip().lower() == target for p in document.paragraphs)


def add_bullets(document: Document, current: Paragraph, items: list[str]) -> Paragraph:
    for item in items:
        style = 'List Bullet' if 'List Bullet' in document.styles else None
        current = insert_paragraph_after(document, current, item, style)
    return current


def update_declaration_date(document: Document, new_date: str) -> None:
    date_pattern = re.compile(r'\b\d{2}/\d{2}/\d{4}\b')
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == '27/10/2025':
            paragraph.text = new_date
            continue

        # Handle declaration lines formatted as "Date: dd/mm/yyyy".
        if text.lower().startswith('date:') and date_pattern.search(text):
            paragraph.text = date_pattern.sub(new_date, paragraph.text)


def patch_programme_framing(document: Document) -> None:
    replacements = {
        'This makes behavioural analytics a suitable topic for a cyber security dissertation.':
            'This makes behavioural analytics a suitable topic for a Computer Science dissertation.',
        'cyber security dissertation': 'Computer Science dissertation',
    }

    for paragraph in document.paragraphs:
        text = paragraph.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            paragraph.text = updated


def inject_chapter2_content(document: Document) -> None:
    probe = '2.6.1 Persona and User Story Synthesis'
    if section_present(document, probe):
        print('SKIP Chapter 2 additions already present')
        return

    insertion_anchor = get_paragraph_before(document, '2.7 Risks and Project Planning')
    current = insertion_anchor

    current = insert_paragraph_after(document, current, '2.6.1 Persona and User Story Synthesis', 'Heading 3')
    current = insert_paragraph_after(
        document,
        current,
        'To strengthen traceability from stakeholder evidence to design decisions, two project personas were consolidated and mapped to implemented user stories.'
    )

    current = insert_paragraph_after(document, current, 'Primary persona: Undergraduate student managing distraction-heavy digital routines.')
    current = add_bullets(document, current, [
        'Goal: reduce unplanned time in high-distraction categories while preserving required communication use.',
        'Pain point: existing usage tools show totals but provide limited intervention guidance.',
        'Success criterion: actionable behaviour-change prompts linked to measurable goals.',
    ])

    current = insert_paragraph_after(document, current, 'Privacy-sensitive persona: reflective user prioritising local control of personal data.')
    current = add_bullets(document, current, [
        'Goal: gain behavioural insight without cloud profiling risk.',
        'Pain point: low trust in opaque retention and telemetry models.',
        'Success criterion: visible consent control, export capability, deletion capability, and pseudonymised identifiers.',
    ])

    current = insert_paragraph_after(document, current, 'Mapped user stories (implemented):')
    current = add_bullets(document, current, [
        'Import messaging and browsing records to establish a behavioural baseline.',
        'Set category and device-aware daily limits.',
        'Create intervention plans to support behaviour interruption when risk increases.',
        'Toggle consent at runtime to immediately stop further collection.',
        'Export and delete personal records on demand.',
    ])

    current = insert_paragraph_after(document, current, '2.6.2 Use Cases and Design Implications', 'Heading 3')
    current = insert_paragraph_after(
        document,
        current,
        'The primary use cases are baseline insight, goal definition, intervention planning, consent management, and data control operations. These use cases directly informed a design that prioritises local-first processing, transparent privacy controls, and intervention-ready analytics rather than passive reporting only.'
    )
    current = insert_paragraph_after(
        document,
        current,
        'A rendered use-case diagram is provided in Appendix J alongside the component, sequence, and activity diagrams.'
    )

    print('INJECTED Chapter 2 additions')


def inject_chapter3_content(document: Document) -> None:
    probe = '3.2.4 Data Model and API Design Rationale'
    if section_present(document, probe):
        print('SKIP Chapter 3 additions already present')
        return

    insertion_anchor = get_paragraph_before(document, '3.3 Using the Artefact and Generating Output')
    current = insertion_anchor

    current = insert_paragraph_after(document, current, '3.2.4 Data Model and API Design Rationale', 'Heading 3')
    current = insert_paragraph_after(
        document,
        current,
        'The implementation uses local SQLite persistence to support an offline-capable, privacy-preserving architecture. Core entities include users, auth tokens, settings, events, summary cache, and habit goals. This structure supports profile isolation, longitudinal event analysis, and goal-oriented intervention planning.'
    )
    current = add_bullets(document, current, [
        'Index strategy supports summary responsiveness across occurred_at, category, source, device, and platform dimensions.',
        'A partial unique index on user_id plus external_id enforces idempotent batch ingestion when upstream systems resend records.',
        'Migration safeguards preserve legacy data while evolving schema fields required by cross-device analytics.',
    ])

    current = insert_paragraph_after(
        document,
        current,
        'API design is organised by behavioural workflow: authentication, consent, event ingestion, imports, summaries, and habit-goal operations. This partitioning supports reproducible evaluation and clear linkage between requirements and endpoint behaviour.'
    )

    current = insert_paragraph_after(document, current, '3.2.5 Privacy/Security Realisation and Architectural Narrative', 'Heading 3')
    current = add_bullets(document, current, [
        'Identifiers are pseudonymised via salted SHA-256 before storage paths are reached.',
        'Strict salt mode prevents insecure production startup with default anonymisation settings.',
        'Consent enforcement is applied in event ingestion routes to ensure user control is operational, not declarative.',
    ])
    current = insert_paragraph_after(
        document,
        current,
        'Rendered component and sequence diagrams are included in Appendix J and correspond to the implemented service boundaries and ingestion-to-summary flow.'
    )

    print('INJECTED Chapter 3 additions')


def inject_chapter4_content(document: Document) -> None:
    probe = '4.1.1 Usability Protocol and Participant Results'
    if section_present(document, probe):
        print('SKIP Chapter 4 additions already present')
        return

    anchor_41 = get_paragraph_before(document, '4.2 System Testing')
    current = anchor_41
    current = insert_paragraph_after(document, current, '4.1.1 Usability Protocol and Participant Results', 'Heading 3')
    current = insert_paragraph_after(
        document,
        current,
        'A structured protocol was applied with three tasks: (1) connect/import and reach dashboard, (2) define a behaviour-change goal, (3) select an intervention action. Participant timing and SUS responses were recorded in reproducible artefacts.'
    )
    current = add_bullets(document, current, [
        'Task timing means (seconds): import 90.0, goal setup 68.5, intervention selection 150.0.',
        'SUS mean: 73.8 out of 100, indicating acceptable baseline usability with refinement opportunities.',
        'Qualitative themes: clearer understanding once intervention options appear; demand for stronger proactive reminders.',
    ])
    current = insert_paragraph_after(
        document,
        current,
        'Full raw usability records and computed summary outputs are provided in Appendix I.'
    )

    anchor_42 = get_paragraph_before(document, '4.3 Critical Review')
    current = anchor_42
    current = insert_paragraph_after(document, current, '4.2.1 System Test Matrix and Outcome Traceability', 'Heading 3')
    current = insert_paragraph_after(
        document,
        current,
        'Automated evidence includes API/integration and end-to-end execution covering authentication, profile isolation, import connectors, habit-goal lifecycle, consent enforcement, idempotent batching, global summary authorisation, and deletion/reset behaviour.'
    )
    current = add_bullets(document, current, [
        'API and integration suite: 14 passed, 0 failed.',
        'End-to-end suite: 1 passed, 0 failed.',
        'Expected versus actual outcomes are consolidated in Appendix H.',
    ])

    anchor_43 = find_paragraph_by_exact_text(document, '4.3 Critical Review')
    current = anchor_43
    current = insert_paragraph_after(document, current, '4.3.1 Evaluation Limits and Improvement Actions', 'Heading 3')
    current = add_bullets(document, current, [
        'Sample size for usability testing remains small (n=2), so conclusions are directional and formative.',
        'Longitudinal relapse behaviour requires multi-week deployment to validate sustained intervention effects.',
        'iOS depth is constrained by project-file completeness in the assessed workspace snapshot.',
    ])
    current = insert_paragraph_after(
        document,
        current,
        'Planned improvements include pre-threshold nudging, richer cross-device intervention triggers, and expanded participant recruitment for stronger generalisability.'
    )

    print('INJECTED Chapter 4 additions')


def inject_appendices(document: Document, figures_dir: Path) -> None:
    probe = 'Appendix G: Interview Prompt Pack'
    if section_present(document, probe):
        print('SKIP Appendix additions already present')
        return

    anchor = find_paragraph_by_exact_text(document, 'Appendix F: Stakeholder Engagement Outcomes Table')
    current = anchor

    current = insert_paragraph_after(document, current, 'Appendix G: Interview Prompt Pack', 'Heading 2')
    current = insert_paragraph_after(
        document,
        current,
        'This appendix records the structured interview prompts used to elicit stakeholder goals, pain points, privacy expectations, and feature priorities.'
    )
    current = insert_paragraph_after(
        document,
        current,
        'Source artefact: myriad-macos-app/artifacts/dissertation/interview-prompts.md'
    )

    current = insert_paragraph_after(document, current, 'Appendix H: Requirements and System Test Traceability', 'Heading 2')
    current = insert_paragraph_after(
        document,
        current,
        'This appendix contains the requirement-to-implementation mapping and expected-versus-actual system test matrix used to support validation claims.'
    )
    current = add_bullets(document, current, [
        'Requirements matrix source: myriad-macos-app/artifacts/dissertation/requirements-traceability.csv',
        'System test matrix source: myriad-macos-app/artifacts/dissertation/system-test-matrix.csv',
    ])

    current = insert_paragraph_after(document, current, 'Appendix I: Usability Raw Data and Computed Summary', 'Heading 2')
    current = add_bullets(document, current, [
        'Raw participant data: myriad-macos-app/artifacts/usability/results-template.csv',
        'Computed summary: myriad-macos-app/artifacts/usability/summary.md',
    ])

    current = insert_paragraph_after(document, current, 'Appendix J: Rendered Design Diagrams', 'Heading 2')
    current = insert_paragraph_after(
        document,
        current,
        'The following rendered figures provide dissertation-formatted architecture artefacts derived from the design model.'
    )

    figures = [
        ('01-e1-use-case-overview.png', 'Figure J1: Use Case Overview'),
        ('02-e2-component-diagram.png', 'Figure J2: Component Diagram'),
        ('03-e3-sequence-diagram-ingestion-to-insight.png', 'Figure J3: Sequence Diagram - Ingestion to Insight'),
        ('04-e4-activity-flow-behavior-change-cycle.png', 'Figure J4: Activity Flow - Behavior-Change Cycle'),
    ]

    for filename, caption in figures:
        image_path = figures_dir / filename
        if not image_path.exists():
            raise FileNotFoundError(f'Missing rendered figure: {image_path}')

        image_paragraph = insert_paragraph_after(document, current, '')
        image_run = image_paragraph.add_run()
        image_run.add_picture(str(image_path), width=Inches(5.8))
        current = insert_paragraph_after(document, image_paragraph, caption)

    print('INJECTED Appendix additions (G-J)')


def main() -> None:
    parser = argparse.ArgumentParser(description='Produce a clean, structured dissertation docx from the original draft.')
    parser.add_argument('--docx', required=True, help='Path to original dissertation docx')
    parser.add_argument('--output', required=True, help='Path to output cleaned docx')
    parser.add_argument(
        '--figures-dir',
        default=None,
        help='Directory containing rendered diagram PNGs. Defaults to artifacts/dissertation/figures in repository.',
    )
    parser.add_argument(
        '--date',
        default=date.today().strftime('%d/%m/%Y'),
        help='Declaration date in dd/mm/YYYY format',
    )
    args = parser.parse_args()

    docx_path = Path(args.docx).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    default_figures = Path(__file__).resolve().parent.parent / 'artifacts' / 'dissertation' / 'figures'
    figures_dir = Path(args.figures_dir).expanduser().resolve() if args.figures_dir else default_figures

    document = Document(str(docx_path))

    update_declaration_date(document, args.date)
    patch_programme_framing(document)
    inject_chapter2_content(document)
    inject_chapter3_content(document)
    inject_chapter4_content(document)
    inject_appendices(document, figures_dir)

    document.save(str(output_path))
    print(f'SAVED {output_path}')


if __name__ == '__main__':
    main()
