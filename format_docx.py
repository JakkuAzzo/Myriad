from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

path = Path('/Users/nathanbrown-bennett/Said Osman/Myriad_Dissertation_Said_Osman.docx')
doc = Document(path)

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

for section in doc.sections:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_separate)
    r_element.append(fld_end)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal_pf = normal.paragraph_format
normal_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
normal_pf.space_after = Pt(0)

for paragraph in doc.paragraphs:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.first_line_indent = None

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pf = paragraph.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                pf.space_after = Pt(0)
                pf.first_line_indent = None

doc.save(path)
print('FORMATTED_OK')
