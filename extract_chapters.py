import sys
from docx import Document

def extract_chapters(docx_path):
    doc = Document(docx_path)
    # The output from previous run shows "Glossary of Terms" as Heading 1 before Chapter 1.
    # We need to correctly identify which Heading 1 corresponds to Chapter 2, 3, 4.
    
    h1_count = 0
    target_chapters = [2, 3, 4] # Based on the logical sequence
    
    # Let's count all Heading 1s and see their text to be sure.
    all_h1s = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 1':
            all_h1s.append((i, para.text.strip()))
    
    # Based on the previous output:
    # Index 34: Glossary of Terms (H1) -> 1st H1
    # Index 36: 1. Introduction and State of the Art Review (H1) -> 2nd H1
    # Index 49: 2. Analysis and Design (H1) -> 3rd H1
    
    # So Chapter 2 starts at the 3rd Heading 1, Chapter 3 is 4th, Chapter 4 is 5th.
    
    current_h1_index = 0
    counters = [0, 0, 0]
    
    for i, para in enumerate(doc.paragraphs):
        sn = para.style.name
        if not sn.startswith('Heading'):
            continue
            
        try:
            level = int(sn.split()[-1])
        except:
            continue
            
        text = para.text.strip()
        
        if level == 1:
            current_h1_index += 1
            counters[1] = 0
            counters[2] = 0
            # If this H1 starts with "2. ", "3. ", or "4. " OR if it's the 3rd, 4th, 5th H1
            # Let's use the explicit "2.", "3.", "4." prefix if available, or sequence.
            # Looking at the output, Chapter 2 starts with "2. ".
            
        # We want Chapters 2, 3, and 4.
        # Let's track the "Chapter Number" based on the Heading 1 text if it starts with a number.
        if level == 1:
            # Try to extract number: "2. Analysis" -> 2
            import re
            match = re.search(r'^(\d+)\.', text)
            if match:
                ch_num = int(match.group(1))
                counters[0] = ch_num
            else:
                # If no number (like Glossary), set to 0 or similar
                counters[0] = -1 

        if counters[0] in [2, 3, 4]:
            print(f"Index: {i} | Heading {level} | {text}")

if __name__ == "__main__":
    extract_chapters("Myriad_Dissertation_Said_Osman_Completed.docx")
